"""文档导出路由 - 支持将 Markdown 内容转换为 DOCX 格式下载"""

import io
import re
from urllib.parse import quote

from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

router = APIRouter(prefix="/api/v1/export", tags=["export"])


class ExportDocxRequest(BaseModel):
    title: str = "未命名文档"
    content: str  # Markdown 格式内容


@router.post("/docx")
async def export_to_docx(req: ExportDocxRequest):
    """将 Markdown 内容转换为 DOCX 文件并返回下载"""
    try:
        doc = Document()

        # 设置文档标题
        doc.add_heading(req.title, level=0)

        # 解析 Markdown 并写入 DOCX
        _markdown_to_docx(doc, req.content)

        # 输出到 BytesIO
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = quote(f"{req.title}.docx")

        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename*=UTF-8''{filename}"
            },
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"导出失败: {str(e)}")


def _markdown_to_docx(doc: Document, markdown_text: str):
    """将 Markdown 文本转换为 DOCX 内容"""
    lines = markdown_text.split("\n")
    i = 0
    in_code_block = False
    code_block_content: list[str] = []
    in_table = False
    table_rows: list[list[str]] = []

    while i < len(lines):
        line = lines[i]

        # 代码块处理
        if line.strip().startswith("```"):
            if in_code_block:
                # 结束代码块
                code_text = "\n".join(code_block_content)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                p.paragraph_format.left_indent = Inches(0.5)
                code_block_content = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_block_content.append(line)
            i += 1
            continue

        # 表格处理
        if line.strip().startswith("|") and line.strip().endswith("|"):
            if not in_table:
                in_table = True
                table_rows = []
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            # 跳过分隔行 (|---|---|)
            if not all(set(c.strip()) <= {"-", ":"} for c in cells):
                table_rows.append(cells)
            i += 1
            continue
        elif in_table:
            # 表格结束，写入
            _write_table(doc, table_rows)
            table_rows = []
            in_table = False
            # 继续处理当前行（不 i += 1）

        # 空行
        if not line.strip():
            i += 1
            continue

        # 标题
        heading_match = re.match(r"^(#{1,6})\s+(.*)", line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            doc.add_heading(_clean_markdown_inline(text), level=min(level, 4))
            i += 1
            continue

        # 无序列表
        list_match = re.match(r"^(\s*)[*\-+]\s+(.*)", line)
        if list_match:
            indent = len(list_match.group(1))
            text = list_match.group(2)
            p = doc.add_paragraph(
                _clean_markdown_inline(text), style="List Bullet"
            )
            if indent >= 2:
                p.paragraph_format.left_indent = Inches(0.25 * (indent // 2))
            i += 1
            continue

        # 有序列表
        ol_match = re.match(r"^(\s*)\d+\.\s+(.*)", line)
        if ol_match:
            indent = len(ol_match.group(1))
            text = ol_match.group(2)
            p = doc.add_paragraph(
                _clean_markdown_inline(text), style="List Number"
            )
            if indent >= 2:
                p.paragraph_format.left_indent = Inches(0.25 * (indent // 2))
            i += 1
            continue

        # 引用块
        if line.strip().startswith(">"):
            text = line.strip().lstrip(">").strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run(_clean_markdown_inline(text))
            run.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)
            i += 1
            continue

        # 分隔线
        if re.match(r"^[-*_]{3,}\s*$", line.strip()):
            doc.add_paragraph("_" * 50)
            i += 1
            continue

        # 普通段落
        p = doc.add_paragraph()
        _add_formatted_text(p, line)
        i += 1

    # 处理未结束的表格
    if in_table and table_rows:
        _write_table(doc, table_rows)


def _write_table(doc: Document, rows: list):
    """写入表格到 DOCX"""
    if not rows:
        return
    num_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            if c_idx < num_cols:
                table.cell(r_idx, c_idx).text = _clean_markdown_inline(
                    cell_text
                )


def _clean_markdown_inline(text: str) -> str:
    """清除 Markdown 内联格式标记，返回纯文本"""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)  # 粗体
    text = re.sub(r"\*(.+?)\*", r"\1", text)  # 斜体
    text = re.sub(r"`(.+?)`", r"\1", text)  # 行内代码
    text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)  # 链接
    text = re.sub(r"~~(.+?)~~", r"\1", text)  # 删除线
    return text.strip()


def _add_formatted_text(paragraph, text: str):
    """解析 Markdown 内联格式，添加带格式的 run"""
    # 按粗体/斜体/行内代码/链接分段
    pattern = r"(\*\*.*?\*\*|\*.*?\*|`.*?`|\[.*?\]\(.*?\)|[^*`\[]+)"
    parts = re.findall(pattern, text)

    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        elif part.startswith("["):
            link_match = re.match(r"\[(.+?)\]\((.+?)\)", part)
            if link_match:
                run = paragraph.add_run(link_match.group(1))
                run.font.color.rgb = RGBColor(74, 158, 255)
                run.underline = True
            else:
                paragraph.add_run(part)
        else:
            paragraph.add_run(part)
